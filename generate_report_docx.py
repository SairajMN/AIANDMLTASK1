from docx import Document
from docx.shared import Inches
import pandas as pd

# Load data for examples
df = pd.read_csv("ai_demo.csv")

doc = Document()
doc.add_heading('Exploratory Data Inspection and Data-Quality Report', 0)

doc.add_heading('1. Title and Purpose', level=1)
doc.add_paragraph('Exploratory Data Analysis (EDA) Report for AI Technologies Dataset')
doc.add_paragraph('This report provides a comprehensive inspection of the dataset to assess its structure, quality, and suitability for machine learning tasks. The dataset contains 10 rows and 4 columns describing AI technologies, with no missing values or duplicates. It is suitable for small-scale ML tasks like classification, but its size limits deep learning applications.')

doc.add_heading('2. Steps and Explanation for Loading Data and Previewing First/Last Rows', level=1)
doc.add_paragraph('To load the data, use pandas\' read_csv function, specifying the file path. Preview the first and last rows to check data consistency and detect potential issues like malformed entries.')
doc.add_paragraph('Steps:')
p = doc.add_paragraph()
p.add_run('• Import pandas as pd.\n').bold = False
p.add_run('• df = pd.read_csv(csv_path)\n').bold = False
p.add_run('• df.head(n) and df.tail(n) for n=5.\n').bold = False
doc.add_paragraph('This reveals all columns are object types, with unique descriptions per row. No obvious formatting issues.')

doc.add_heading('3. Manual + Programmatic Guidance to Classify Features into Numeric, Categorical, Ordinal, and Binary', level=1)
doc.add_paragraph('Inspect dtypes and unique counts. All columns are object; no numeric or binary. Technology and Description are high-cardinality (unique=10), Year Introduced categorical (8 unique), Company categorical (9 unique). No ordinal evident without domain knowledge.')

doc.add_heading('4. Use df.info() and df.describe() Style Outputs — Explain Exactly What to Look for and How to Interpret Mean/Median/Std/Skew/Outliers', level=1)
doc.add_paragraph('df.info() provides a concise summary of the DataFrame\'s structure. Key elements:')
doc.add_paragraph('• Class and type: Confirms it\'s a pandas DataFrame.')
doc.add_paragraph('• RangeIndex: Number of rows (10 entries) and index range (0 to 9).')
doc.add_paragraph('• Data columns: Total columns (4), with details per column:')
doc.add_paragraph('  - Column name.')
doc.add_paragraph('  - Non-Null Count: Number of non-missing values (10 for all, indicating no missing data).')
doc.add_paragraph('  - Dtype: Data type (object for all, meaning strings/text).')
doc.add_paragraph('• dtypes summary: Counts of each data type (all object).')
doc.add_paragraph('• Memory usage: Total memory used (452.0+ bytes, low due to small size).')
doc.add_paragraph('Look for inconsistencies like unexpected dtypes (e.g., numeric as object), missing values (non-null < total rows), or high memory usage indicating large data.')
doc.add_paragraph('df.describe(include=\'all\') summarizes statistics for all columns. For numeric columns (none here), it includes count, mean, std, min, 25%, 50% (median), 75%, max.')
doc.add_paragraph('Interpretation of numeric stats:')
doc.add_paragraph('• Mean: Average value; compare to median to detect skew (mean > median = right skew, mean < median = left skew).')
doc.add_paragraph('• Std (standard deviation): Measures spread; high std indicates variability.')
doc.add_paragraph('• Min/Max: Range; extreme values may indicate outliers.')
doc.add_paragraph('• Quartiles (25%, 50%, 75%): Median splits data; IQR = Q3 - Q1 for outlier detection (values outside Q1-1.5*IQR or Q3+1.5*IQR are potential outliers).')
doc.add_paragraph('• Skew: Not directly shown, but infer from mean/median difference.')
doc.add_paragraph('For categorical columns (all here):')
doc.add_paragraph('• Count: Total non-null entries (10).')
doc.add_paragraph('• Unique: Number of distinct values (10 for Technology/Description, indicating all unique; 8 for Year Introduced, 9 for Company).')
doc.add_paragraph('• Top: Most frequent value (e.g., "Artificial Intelligence" for Technology, "2010s" for Year Introduced, "General" for Company).')
doc.add_paragraph('• Freq: Frequency of top value (1 for most, 2 for 2010s and General).')
doc.add_paragraph('Interpretation: High unique relative to count suggests diversity; low freq for top indicates balance, high freq suggests mode dominance. Here, balanced except slight mode in Year Introduced and Company.')
doc.add_paragraph('df.info() output:')
doc.add_paragraph('<class \'pandas.core.frame.DataFrame\'>\nRangeIndex: 10 entries, 0 to 9\nData columns (total 4 columns):\n #   Column           Non-Null Count  Dtype \n---  ------           --------------  ----- \n 0   Technology       10 non-null     object\n 1   Description      10 non-null     object\n 2   Year Introduced  10 non-null     object\n 3   Company          10 non-null     object\ndtypes: object(4)\nmemory usage: 452.0+ bytes')
doc.add_paragraph('df.describe() output:')
doc.add_paragraph('                     Technology                                                                                                                       Description Year Introduced  Company\ncount                        10                                                                                                                                10              10       10\nunique                       10                                                                                                                                10               8        9\ntop     Artificial Intelligence  A broad field of computer science focused on creating machines that can perform tasks that typically require human intelligence.           2010s  General\nfreq                          1                                                                                                                                 1               2        2')

doc.add_heading('5. Display Unique Values for Categorical Columns and Explain How to Handle Inconsistent Labels and High-Cardinality Columns', level=1)
doc.add_paragraph('Unique values for Year Introduced: 2010s (2), 1960s (2), 1956 (1), 1959 (1), 1950s (1), 2012 (1), 1970s (1), 1940s (1). Consistent. For Company: General (2), Various (1), Google (1), IBM (1), MIT (1), DeepMind (1), OpenAI (1), Stanford (1), Warren McCulloch (1). No inconsistencies. High-cardinality in Technology/Description: Group or use embeddings.')

doc.add_heading('6. Identify and Evaluate the Target Variable: Counts, Proportions, and Imbalance Ratio; Give Thresholds for Concern and Recommended Fixes (SMOTE, Class Weights, Undersample)', level=1)
doc.add_paragraph('Target: Company. Counts: General 2, others 1 each. Proportions: General 0.2, others 0.1. Ratio: 2.00 (low imbalance). Threshold: >5 moderate, >10 severe. No fixes needed; minor imbalance acceptable.')

doc.add_heading('7. Analyze Dataset Size and Suitability for Classical ML vs Deep Learning with Clear Heuristics', level=1)
doc.add_paragraph('Size: 10 rows, 4 columns. Heuristics: Too small for deep learning (<1000 rows); suitable for classical ML like Random Forest for small datasets.')

doc.add_heading('8. Deeply Explain Data Quality Issues (Missing Values Types MCAR/MAR/MNAR, Duplicates, Outliers, Leakage), Suggested Remediation Tactics per Issue, and Code Examples', level=1)
doc.add_paragraph('No missing values (MCAR/MAR/MNAR not applicable). No duplicates. No numeric columns, so no outliers. No leakage evident. Quality excellent; no remediation needed.')

doc.add_heading('9. Provide Ready-to-Run Pandas Code', level=1)
code_block = '''```python
import pandas as pd
import numpy as np

csv_path = "ai_demo.csv"
target_column = "Company"

df = pd.read_csv(csv_path)

print("First 5 rows:")
print(df.head().to_string(index=False))
print("\\nLast 5 rows:")
print(df.tail().to_string(index=False))

print("\\ndf.info():")
df.info()
print("\\ndf.describe():")
print(df.describe(include='all').to_string())

summary = pd.DataFrame({
    'dtype': df.dtypes,
    'unique': df.nunique(),
    'missing%': (df.isna().sum() / len(df)) * 100
})
print("\\nColumn Summary:")
print(summary.to_string())

def suggest_type(col):
    if df[col].dtype in ['int64', 'float64']:
        return 'numeric'
    elif df[col].nunique() == 2:
        return 'binary'
    elif df[col].nunique() < 10:
        return 'categorical'
    else:
        return 'high-cardinality'
summary['suggested_type'] = summary.index.map(suggest_type)
print("\\nSuggested Types:")
print(summary[['suggested_type']].to_string())

categoricals = [col for col in df.columns if df[col].dtype == 'object' or df[col].nunique() < 20]
for col in categoricals:
    print(f"\\nValue counts for {col}:")
    print(df[col].value_counts().to_string())

print(f"\\nTarget ({target_column}) counts:")
print(df[target_column].value_counts().to_string())
print(f"\\nTarget proportions:")
print(df[target_column].value_counts(normalize=True).to_string())
imbalance_ratio = df[target_column].value_counts().max() / df[target_column].value_counts().min()
print(f"Imbalance ratio: {imbalance_ratio:.2f}")

missing_table = df.isna().sum().reset_index()
missing_table.columns = ['column', 'missing_count']
print("\\nMissingness Table:")
print(missing_table.to_string())

print(f"\\nDuplicate rows: {df.duplicated().sum()}")

numeric_cols = df.select_dtypes(include=[np.number]).columns
if len(numeric_cols) > 1:
    corr = df[numeric_cols].corr()
    print("\\nCorrelation matrix:")
    print(corr.to_string())

for col in numeric_cols:
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    outliers = ((df[col] < Q1 - 1.5 * IQR) | (df[col] > Q3 + 1.5 * IQR)).sum()
    print(f"Outliers in {col}: {outliers}")

summary.to_csv('column_summary.csv')
print("\\nSaved column_summary.csv")
```'''
doc.add_paragraph(code_block)

doc.add_heading('10. Short Checklist and Recommended Next Steps for Modeling and Concise Summary of Findings', level=1)
doc.add_paragraph('Checklist:')
p = doc.add_paragraph()
p.add_run('• [ ] Encode categoricals (one-hot for Company).\n').bold = False
p.add_run('• [ ] Feature engineer Description (TF-IDF).\n').bold = False
p.add_run('• [ ] Split data; train simple model.\n').bold = False
p.add_run('• [ ] Evaluate with cross-validation.\n').bold = False
doc.add_paragraph('Next Steps: Address high-cardinality; consider NLP for text features.')
doc.add_paragraph('Summary: Small, clean dataset (10 rows, 4 columns); all object types; target Company mildly imbalanced (ratio 2.0); no issues; ideal for classical ML with text preprocessing.')

doc.save('eda_report.docx')
print("DOCX report saved as eda_report.docx")
