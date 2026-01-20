
import pandas as pd
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ---------- CONFIG ----------
FILEPATH = "ai_demo.csv"   # <- change to your dataset path (or .xlsx/.json etc.)
N_PREVIEW = 5                # how many rows for head/tail
READ_KWARGS = {}             # extra kwargs for pd.read_csv, e.g. encoding, sep, decimal, thousands
# Example: READ_KWARGS = {"encoding": "utf-8", "sep": ",", "decimal": ".", "thousands": ","}
# ----------------------------

def load_dataset(path: str, read_kwargs: dict):
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")
    suffix = p.suffix.lower()
    if suffix in {".csv", ".txt"}:
        df = pd.read_csv(path, **read_kwargs)
    elif suffix in {".xls", ".xlsx"}:
        df = pd.read_excel(path, **read_kwargs)
    elif suffix in {".json"}:
        df = pd.read_json(path, **read_kwargs)
    elif suffix in {".parquet"}:
        df = pd.read_parquet(path, **read_kwargs)
    else:
        # fallback try CSV
        df = pd.read_csv(path, **read_kwargs)
    return df

def print_overview(df: pd.DataFrame, n: int = 5):
    print("1) SHAPE (rows, columns):")
    print(df.shape)
    print("\n2) FIRST", n, "ROWS (head):")
    # to_string(index=False) prints a simple table, suitable for Word paste
    print(df.head(n).to_string(index=False))
    print("\n3) LAST", n, "ROWS (tail):")
    print(df.tail(n).to_string(index=False))
    print("\n4) COLUMN INFO (dtypes & non-null counts):")
    print(df.info(verbose=True, show_counts=True))
    print("\n5) MISSING VALUES per column:")
    print(df.isna().sum().to_string())
    print("\n6) QUICK STATISTICS:")
    print(df.describe().to_string())
    print("\n7) MEMORY USAGE (bytes):")
    print(df.memory_usage(deep=True).sum())

def save_head_tail_to_docx(df: pd.DataFrame, filename="preview_head_tail.docx", n=5):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
    doc = Document()
    doc.add_heading("Dataset preview", level=1)
    doc.add_paragraph(f"Shape: {df.shape[0]} rows × {df.shape[1]} columns")
    doc.add_heading("First rows (head)", level=2)
    # Add as preformatted text to preserve spacing
    doc.add_paragraph(df.head(n).to_string(index=False))
    doc.add_heading("Last rows (tail)", level=2)
    doc.add_paragraph(df.tail(n).to_string(index=False))
    doc.add_heading("Column types & non-null counts", level=2)
    # info() does not return a string; construct a basic text summary
    col_summaries = []
    for col in df.columns:
        col_summaries.append(f"{col} — dtype: {df[col].dtype}, non-null: {df[col].notna().sum()}")
    doc.add_paragraph("\n".join(col_summaries))
    doc.save(filename)
    return filename

if __name__ == "__main__":
    # 1) load
    df = load_dataset(FILEPATH, READ_KWARGS)

    # 2) print a clean overview suitable for copying into Word
    print_overview(df, N_PREVIEW)

    # 3) optional: write a Word file with the head & tail
    if DOCX_AVAILABLE:
        out_doc = save_head_tail_to_docx(df, filename="dataset_preview.docx", n=N_PREVIEW)
        print(f"\nSaved Word preview to: {out_doc}")
    else:
        print("\npython-docx not available. To save to .docx install: pip install python-docx")
